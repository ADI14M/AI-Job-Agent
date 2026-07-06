import pdfplumber
import docx
from app.core.logger import system_logger
from app.utils.file_utils import get_file_extension
from app.utils.text_utils import clean_text

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(layout=True)
                if page_text:
                    text += page_text + "\n"
                
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        for row in table:
                            text += " | ".join([str(cell) for cell in row if cell]) + "\n"
    except Exception as e:
        system_logger.error(f"Error extracting text from PDF {file_path}: {e}")
        raise ValueError(f"Failed to read PDF: {e}")
    return clean_text(text)

def extract_text_from_docx(file_path: str) -> str:
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
            
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                text += row_text + "\n"
    except Exception as e:
        system_logger.error(f"Error extracting text from DOCX {file_path}: {e}")
        raise ValueError(f"Failed to read DOCX: {e}")
    return clean_text(text)

def extract_text(file_path: str) -> str:
    ext = get_file_extension(file_path)
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")
